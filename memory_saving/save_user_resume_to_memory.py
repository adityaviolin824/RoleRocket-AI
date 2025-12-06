import json
from typing import Dict, Any, List
from pathlib import Path

import numpy as np
import pdfplumber
import docx
# import easyocr
from PIL import Image
from pydantic import BaseModel
from agents import Agent, Runner, trace
from agents.mcp import MCPServerStdio
from dotenv import load_dotenv
import zipfile
import xml.etree.ElementTree as ET

from memory_saving.memory_mcp_config import MCP_PARAMS, ensure_memory_dir

load_dotenv(override=True)

_ocr_reader = None

##### UNCOMMENT IF YOU WANT TO USE IMAGE RESUMES, COMMENTED DUE TO RENDER FREE TIER CPU LIMITATIONS
# def get_ocr_reader() -> easyocr.Reader:
#     """Lazy initialize EasyOCR reader."""
#     global _ocr_reader
#     if _ocr_reader is None:
#         print("[OCR] Initializing EasyOCR...")
#         _ocr_reader = easyocr.Reader(["en"], gpu=False)
#         print("[OCR] EasyOCR ready!")
#     return _ocr_reader
#
#
# def ocr_image_file(path: Path) -> str:
#     reader = get_ocr_reader()
#     arr = reader.readtext(str(path), detail=0, paragraph=True)
#     return "\n".join(arr)
#
#
# def ocr_image_pil(img: Image.Image) -> str:
#     reader = get_ocr_reader()
#     img_array = np.array(img)
#     arr = reader.readtext(img_array, detail=0, paragraph=True)
#     return "\n".join(arr)

# =========================
# Experience summary (companies + roles only)
# =========================
def detect_role_type(title: str, company: str, responsibilities: List[str]) -> str:
    """
    Heuristic classification: 'internship', 'research', or 'full_time'.
    This is kept in case we want to use it later,
    but we do NOT compute durations or intervals.
    """
    text = " ".join([str(title or ""), str(company or "")] + (responsibilities or [])).lower()
    if "intern" in text:
        return "internship"
    if "research" in text or "researcher" in text or "fellow" in text:
        return "research"
    return "full_time"


def compute_experience_summary(parsed_resume: dict) -> dict:
    """
    Takes the LLM parsed resume dict and computes:
      - experience_detail_computed: per role info (no dates or durations computed)
      - companies: unique list of companies in first-seen order
      - roles: unique list of role titles in first-seen order

    We explicitly DO NOT compute:
      - exact intervals
      - non-overlapping months
      - merged intervals
      - total experience months/years (user will input these separately)
    """
    roles = parsed_resume.get("experience", []) or []
    per_role: List[Dict[str, Any]] = []

    companies: List[str] = []
    roles_list: List[str] = []
    seen_companies = set()
    seen_roles = set()

    for item in roles:
        title = (item.get("title") or "").strip()
        company = (item.get("company") or "").strip()
        resps = item.get("responsibilities", []) or []
        # normalize responsibilities to list of strings
        if isinstance(resps, str):
            resps = [resps]
        start_raw = (item.get("start_date") or "").strip()
        end_raw = (item.get("end_date") or "").strip()
        current_flag = bool(item.get("current", False))

        rtype = detect_role_type(title, company, resps)

        per_role.append(
            {
                "title": title,
                "company": company,
                "start_date": start_raw,
                "end_date": end_raw,
                "current": current_flag,
                "type": rtype,
            }
        )

        if company and company not in seen_companies:
            seen_companies.add(company)
            companies.append(company)

        if title and title not in seen_roles:
            seen_roles.add(title)
            roles_list.append(title)

    experience_summary = {
        "experience_detail_computed": per_role,
        "companies": companies,
        "roles": roles_list,
    }
    return experience_summary


# =========================
# CONTRIBUTIONS aggregation
# =========================
def compute_contributions(parsed_resume: dict) -> List[str]:
    """
    Build a contributions list by combining:
      - company contributions: flatten responsibilities from experience items
      - personal projects: from keys like 'projects', 'personal_projects', 'side_projects'
      - existing 'contributions' field if already present
    Returns a list of contribution strings.
    """
    contributions: List[str] = []

    # 1) Existing contributions provided by parser (preserve if present)
    existing = parsed_resume.get("contributions", []) or []
    if isinstance(existing, str):
        if existing.strip():
            contributions.append(existing.strip())
    elif isinstance(existing, list):
        for c in existing:
            if isinstance(c, str) and c.strip():
                contributions.append(c.strip())

    # 2) Company contributions from experience responsibilities
    exp_items = parsed_resume.get("experience", []) or []
    for item in exp_items:
        title = (item.get("title") or "").strip()
        company = (item.get("company") or "").strip()
        resps = item.get("responsibilities", []) or []
        if isinstance(resps, str):
            resps = [resps]
        for r in resps:
            if not r or not str(r).strip():
                continue
            # Format: "Company - Title: responsibility"
            prefix_parts = []
            if company:
                prefix_parts.append(company)
            if title:
                prefix_parts.append(title)
            prefix = " - ".join(prefix_parts) if prefix_parts else "Company"
            contribution_line = f"{prefix}: {str(r).strip()}"
            contributions.append(contribution_line)

    # 3) Personal / side projects
    # Accept multiple possible keys for flexibility
    project_keys = ["projects", "personal_projects", "side_projects"]
    for key in project_keys:
        projects = parsed_resume.get(key, []) or []
        if isinstance(projects, str):
            if projects.strip():
                contributions.append(f"Project: {projects.strip()}")
        elif isinstance(projects, list):
            for p in projects:
                if not p:
                    continue
                if isinstance(p, str):
                    if p.strip():
                        contributions.append(f"Project: {p.strip()}")
                elif isinstance(p, dict):
                    title = (p.get("title") or p.get("name") or "").strip()
                    desc = (p.get("description") or p.get("summary") or p.get("details") or "").strip()
                    if title and desc:
                        contributions.append(f"Project: {title} - {desc}")
                    elif title:
                        contributions.append(f"Project: {title}")
                    elif desc:
                        contributions.append(f"Project: {desc}")

    # 4) Publications, open source, etc. (optional)
    pub_keys = ["publications", "open_source", "oss"]
    for key in pub_keys:
        items = parsed_resume.get(key, []) or []
        if isinstance(items, str):
            if items.strip():
                contributions.append(items.strip())
        elif isinstance(items, list):
            for it in items:
                if isinstance(it, str) and it.strip():
                    contributions.append(it.strip())
                elif isinstance(it, dict):
                    title = (it.get("title") or it.get("name") or "").strip()
                    if title:
                        contributions.append(title)

    # Deduplicate while preserving order
    seen = set()
    deduped: List[str] = []
    for c in contributions:
        key = c.strip()
        if key and key not in seen:
            seen.add(key)
            deduped.append(key)

    return deduped


# =========================
# BULLETPROOF PDF EXTRACTOR
# =========================
def extract_text_from_pdf(path: Path) -> str:
    """
    BULLETPROOF PDF extraction with multiple fallbacks.
    OCR based fallbacks are disabled in this deployment.
    """
    try:
        print(f"[PDF] Processing {path.name}")
        text_chunks: List[str] = []
        page_stats = {"text_pages": 0, "ocr_pages": 0, "total_pages": 0}

        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_stats["total_pages"] += 1
                print(f"[PDF] Page {page_num}/{len(pdf.pages)}")

                # 1) Native text
                native_text = page.extract_text() or ""
                if native_text.strip():
                    text_chunks.append(native_text.strip())
                    page_stats["text_pages"] += 1
                    print(f"[PDF] Page {page_num}: Native ({len(native_text)} chars)")
                    continue

                # 2) Tables fallback (no OCR)
                tables = page.extract_tables()
                table_text = []
                for table_num, table in enumerate(tables or []):
                    if table:
                        for row in table:
                            row_text = [str(cell or "") for cell in row if cell]
                            if row_text:
                                table_text.append(" | ".join(row_text))
                if table_text:
                    text_chunks.append(f"[PAGE {page_num} TABLES]\n" + "\n".join(table_text))
                    # page_stats["ocr_pages"] += 1  # would have been OCR tables, but OCR is disabled
                    continue

                # 3) OCR based fallbacks are disabled
                # try:
                #     page_img = page.to_image(resolution=300)
                #     ocr_text = ocr_image_pil(page_img.original)
                #     if ocr_text.strip():
                #         text_chunks.append(f"[PAGE {page_num} OCR]\n{ocr_text.strip()}")
                #         page_stats["ocr_pages"] += 1
                #         print(f"[PDF] Page {page_num}: OCR ({len(ocr_text)} chars)")
                #         continue
                # except Exception as e:
                #     print(f"[PDF] Page {page_num} OCR failed: {e}")
                #
                # try:
                #     full_img = page.to_image(resolution=400)
                #     full_ocr = ocr_image_pil(full_img.original)
                #     if full_ocr.strip():
                #         text_chunks.append(f"[PAGE {page_num} FULL OCR]\n{full_ocr}")
                #         page_stats["ocr_pages"] += 1
                #         print(f"[PDF] Page {page_num}: Full OCR ({len(full_ocr)} chars)")
                # except Exception as e:
                #     print(f"[PDF] Page {page_num} Full OCR failed: {e}")

        print(
            f"[PDF] SUMMARY: {page_stats['text_pages']} text + "
            f"{page_stats['ocr_pages']} OCR / {page_stats['total_pages']} pages"
        )

        result = "\n\n".join(text_chunks)
        if not result.strip():
            raise_extraction_error(path)

        print(f"[PDF] TOTAL: {len(result)} chars")
        return result

    except Exception as e:
        print(f"[PDF] CRITICAL ERROR: {e}")
        raise ValueError(f"[PDF] Failed {path.name}: {str(e)}")


def raise_extraction_error(path: Path):
    raise ValueError(f"[PDF] No text from {path.name}. Image-only/corrupted?")


# =========================
# BULLETPROOF DOCX EXTRACTOR
# =========================
def extract_text_from_docx(path: Path) -> str:
    """Enhanced DOCX extraction (paragraphs, tables, XML fallback)."""
    try:
        doc = docx.Document(str(path))
        full_text: List[str] = []
        print(f"[DOCX] Processing {path.name}")

        para_count = 0
        for p in doc.paragraphs:
            para_text = [run.text.strip() for run in p.runs if run.text.strip()]
            if para_text:
                full_text.append(" ".join(para_text))
                para_count += 1

        table_count = 0
        for table in doc.tables:
            for row in table.rows:
                row_text: List[str] = []
                for cell in row.cells:
                    for p in cell.paragraphs:
                        cell_text = [run.text.strip() for run in p.runs if run.text.strip()]
                        row_text.extend(cell_text)
                if row_text:
                    full_text.append(" | ".join(row_text))
                    table_count += 1

        print(f"[DOCX] Extracted: {para_count} paras, {table_count} tables")

        xml_text = extract_text_from_docx_xml(path)
        if xml_text.strip():
            full_text.append(xml_text)
            print(f"[DOCX] XML added: {len(xml_text)} chars")

        result = "\n".join(full_text)
        print(f"[DOCX] TOTAL: {len(result)} chars")
        return result if result.strip() else raise_extraction_error(path)

    except Exception as e:
        print(f"[DOCX] ERROR: {e}")
        raise ValueError(f"[DOCX] Failed: {e}")


def extract_text_from_docx_xml(path: Path) -> str:
    """XML parsing for DOCX textboxes and shapes."""
    try:
        with zipfile.ZipFile(str(path), "r") as zip_ref:
            if "word/document.xml" in zip_ref.namelist():
                with zip_ref.open("word/document.xml") as xml_file:
                    tree = ET.parse(xml_file)
                    root = tree.getroot()
                    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    text_elements = root.findall(".//w:t", ns)
                    return "\n".join(
                        [elem.text for elem in text_elements if elem.text and elem.text.strip()]
                    )
    except Exception:
        pass
    return ""


# =========================
# UNIVERSAL EXTRACTOR
# =========================
def extract_resume_text(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    elif ext in [".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"]:
        # OCR is disabled - image resumes are not supported in this deployment
        raise ValueError(
            f"Image resume formats like {ext} are not supported in this environment. "
            f"Please upload a PDF or DOCX file instead."
        )
    else:
        raise ValueError(f"Unsupported: {ext}. Use PDF, DOCX, PNG, or JPG")


# =========================
# Pydantic models
# =========================
class BasicInfo(BaseModel):
    full_name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    headline: str = ""


class ExperienceItem(BaseModel):
    title: str = ""
    company: str = ""
    location: str = ""
    start_date: str = ""
    end_date: str = ""
    current: bool = False
    responsibilities: List[str] = []


class EducationItem(BaseModel):
    degree: str = ""
    field: str = ""
    institution: str = ""
    start_date: str = ""
    end_date: str = ""


class Skills(BaseModel):
    technical: List[str] = []
    soft: List[str] = []


class Resume(BaseModel):
    basic_info: BasicInfo
    summary: str
    skills: Skills
    experience: List[ExperienceItem]
    education: List[EducationItem]
    years_experience: float


# =========================
# Resume parser agent
# =========================
def build_resume_parser_agent(model: str = "gpt-4o-mini") -> Agent[Resume]:
    instructions = (
        "You are a strict resume parser.\n"
        "Extract structured data from raw resume text into the provided Resume schema.\n"
        "Use empty strings for missing text fields, empty lists for missing arrays, "
        "false for missing booleans, and 0 for missing numeric fields. "
        "Estimate total years of experience as a number. "
        "For dates, use 'YYYY-MM' when possible, otherwise 'YYYY'. "
        "Support quarter formats like 'Q2 2024' and season formats like 'Summer 2023'. "
        "Do not invent jobs or degrees that are not clearly present in the text."
    )
    agent = Agent[Resume](
        name="resume_parser",
        instructions=instructions,
        model=model,
        output_type=Resume,
    )
    return agent


async def parse_resume_with_llm(raw_text: str, agent: Agent[Resume]) -> Dict[str, Any]:
    prompt = (
        "Parse the following resume text into the structured Resume schema. "
        "Follow your instructions and fill all fields consistently.\n\nRESUME:\n"
        f"{raw_text}"
    )
    try:
        with trace("Parsing and Saving Resume"):
            result = await Runner.run(agent, prompt)
            resume_obj: Resume = result.final_output
            return resume_obj.model_dump()
    except Exception as e:
        raise Exception(f"Error parsing resume with Agent: {str(e)}")


# =========================
# Memory saving
# =========================
async def save_resume_profile_to_memory(parsed_resume: dict) -> None:
    ensure_memory_dir()
    params = MCP_PARAMS

    basic = parsed_resume.get("basic_info", {})
    skills = parsed_resume.get("skills", {})
    summary_text = parsed_resume.get("summary", "").strip()
    exp_summary = parsed_resume.get("experience_summary", {})

    companies = exp_summary.get("companies", [])
    roles = exp_summary.get("roles", [])

    contributions = parsed_resume.get("contributions", []) or []
    # top contributions preview
    top_contributions = contributions[:10]

    payload = {
        "type": "resume_profile",
        "full_name": basic.get("full_name", ""),
        "email": basic.get("email", ""),
        "headline": basic.get("headline", ""),
        "location": basic.get("location", ""),
        # We do NOT compute experience; we use whatever the parser gave,
        "years_experience": parsed_resume.get("years_experience", 0.0),
        "top_technical_skills": skills.get("technical", [])[:15],
        "summary": summary_text,
        "companies": companies,
        "roles": roles,
        "contributions": contributions,
        "top_contributions": top_contributions,
    }

    async with MCPServerStdio(params=params, client_session_timeout_seconds=120) as mcp_server:
        try:
            await mcp_server.call_tool("delete_entity", {"name": "resume_profile"})
            print("[Memory] Deleted existing 'resume_profile' entity (if it existed).")
        except Exception as e:
            print(f"[Memory] Warning while deleting old entity: {e}")

        entities_arg = {
            "entities": [
                {
                    "name": "resume_profile",
                    "entityType": "resume_profile",
                    "observations": [json.dumps(payload, indent=2)],
                }
            ]
        }

        result = await mcp_server.call_tool("create_entities", entities_arg)

        print("[Memory] Saved resume profile to memory.")
        print("[Memory] Tool result:")
        print(result)


# =========================
# Orchestrator
# =========================
async def pipeline_process_resume_file(
    path: Path,
    save_to_memory: bool = False,
    model: str = "gpt-4o-mini",
) -> None:
    resume_agent = build_resume_parser_agent(model=model)
    raw_text = extract_resume_text(path)
    parsed_resume = await parse_resume_with_llm(raw_text, resume_agent)

    # Compute experience summary from parsed roles (companies + roles only)
    try:
        experience_summary = compute_experience_summary(parsed_resume)
        parsed_resume["experience_summary"] = experience_summary
    except Exception as e:
        print(f"[ResumeCompute] Warning computing experience summary: {e}")
        parsed_resume["experience_summary"] = {}

    # Compute aggregated contributions (company contributions + personal projects)
    try:
        contributions = compute_contributions(parsed_resume)
        parsed_resume["contributions"] = contributions
    except Exception as e:
        print(f"[ResumeCompute] Warning computing contributions: {e}")
        parsed_resume["contributions"] = []

    basic = parsed_resume.get("basic_info", {})
    skills = parsed_resume.get("skills", {})
    summary_text = parsed_resume.get("summary", "").strip() or "No summary detected."
    top_skills = skills.get("technical", [])[:10]
    top_skills_str = ", ".join(top_skills) if top_skills else "None detected"

    exp_sum = parsed_resume.get("experience_summary", {})
    detail = exp_sum.get("experience_detail_computed", [])
    num_roles = len(detail)
    companies = exp_sum.get("companies", []) or []
    roles_list = exp_sum.get("roles", []) or []

    companies_str = ", ".join(companies) if companies else "None detected"
    roles_str = ", ".join(roles_list) if roles_list else "None detected"

    yrs_llm = parsed_resume.get("years_experience", "N/A")

    print("\n================= RESUME SNAPSHOT =================\n")
    print(f"Name:        {basic.get('full_name', '')}")
    print(f"Email:       {basic.get('email', '')}")
    print(f"Headline:    {basic.get('headline', '')}")
    print(f"Location:    {basic.get('location', '')}")
    print(f"Top Skills:  {top_skills_str}")
    print(f"Years Exp (LLM estimate):    {yrs_llm}")
    print(f"Parsed roles count:          {num_roles}")
    print(f"Companies:                   {companies_str}")
    print(f"Roles:                       {roles_str}")

    contributions = parsed_resume.get("contributions", []) or []
    contrib_count = len(contributions)
    print("\n---------------- CONTRIBUTIONS ----------------\n")
    print(f"Total contributions detected: {contrib_count}")
    for i, c in enumerate(contributions[:6], start=1):
        print(f"{i}. {c}")
    if contrib_count > 6:
        print(f"... and {contrib_count - 6} more contributions (saved to memory).")

    print("\n------------------ SUMMARY CHECK ------------------\n")
    print(summary_text)
    print("\nThese are the details extracted, they will be saved for later steps.")
    print("\n===================================================\n")

    if save_to_memory:
        await save_resume_profile_to_memory(parsed_resume)
    else:
        print("Skipping memory save.")
